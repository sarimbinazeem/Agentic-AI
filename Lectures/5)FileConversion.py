import re # for clean text

from docx import Document
from pypdf import PdfReader

#for directories
from pathlib import Path

RAW_DIR = Path("data/raw")
DOCS_DIR = Path("docs")


def readText(source_path):
    return source_path.read_text(encoding="utf-8")

def readPDF(source_path):
    pdf = PdfReader(str(source_path))
    pages = []
    
    for i,page in enumerate(pdf.pages,start=1): #enumarate gives index as well as pages
        page_text = page.extract_text() or ""
        if page_text.strip():
            pages.append(page_text.strip())
            
    return "\n\n".join(pages) #this joins double line to every line of pages list

def readDocx(source_path):
    doc = Document(source_path)
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()] #this return those paragraphs that are not empty
    return "\n\n".join(parts)

def cleanText(text):
    text = text.replace("\r\n","\n").replace("\r","\n") #replace spaces and give new line after each word
    text = re.sub(r"[ \t]+"," ",text) #replace ONE OR MORE SPACES OR TABS with single space in TEXT
    text = re.sub(r"\n{3,}","\n\n",text) #replace more than 2 new lines with 2 NEW LINES in text
    return text.strip()

    
    
def convertOne(source_path,output_dir):

    suffix = source_path.suffix.lower()
    output_path = output_dir / f"{source_path.stem}.txt"
    
    if suffix == ".txt":
        text = readText(source_path)
        
    elif suffix == ".pdf":
        text = readPDF(source_path)
        
    elif suffix == ".docx":
        text = readDocx(source_path)
        
    else:
        print(f"Unsupported File: {source_path.name}")
        return None
    
    output_dir.mkdir(parents=True,exist_ok =True)
    output_path.write_text(cleanText(text),encoding="utf-8")
    return output_path

def ensure_sample_docx():
    path = RAW_DIR / "sample_company_policy.docx"
    if path.exists():
        return path 
    
    RAW_DIR.mkdir(parents=True,exist_ok=True)
    
    doc = Document()
    doc.add_heading("Northstar Tech — Remote Work Policy", level=1)
    doc.add_paragraph(
        "Effective date: January 2025. All full-time employees may work remotely "
        "up to 3 days per week with manager approval."
    )
    doc.add_heading("Equipment", level=2)
    doc.add_paragraph(
        "The company provides a laptop and monitor for home office setup. "
        "Employees must use company-approved VPN when accessing internal systems."
    )
    doc.add_heading("Core Hours", level=2)
    doc.add_paragraph(
        "Employees must be available between 10:00 AM and 3:00 PM in their local timezone."
    )
    doc.add_heading("Expense Reimbursement", level=2)
    doc.add_paragraph(
        "Internet stipend of $50 per month is available for eligible remote employees."
    )
    doc.save(path)
    print(f"Created sample DOCX: {path}")
    return path
    
    
def main():
    print("=== File Conversion: PDF / DOCX / TXT -> docs/*.txt ===\n")
    RAW_DIR.mkdir(parents=True,exist_ok=True)
    DOCS_DIR.mkdir(parents=True,exist_ok=True)
    ensure_sample_docx() #create if there is no sample document
    
    #create list of pdf files AND FIND ALL THROUGH GLOB
    pdfs = list(RAW_DIR.glob("*.pdf"))
    #if there are no pdfs
    if not pdfs:
        print("No PDF in data/raw/. Add data/raw/sample_document.pdf to demo PDF extraction.")
        print("Continuing with TXT and DOCX only.\n")
        
    #sort the sources gathered
    
    sources = sorted(
        list(RAW_DIR.glob("*.txt")) +
        list(RAW_DIR.glob("*.pdf")) +
        list(RAW_DIR.glob("*.docx"))
    )
    
    if not sources:
        print("No source files found in data/raw/")
        return
    
    print(f"Found {len(sources)} file(s):")
    for s in sources:
        print(f"-{s.name}")
        
    print("\nConverting...")
    
    for source in sources:
        output = convertOne(source,DOCS_DIR)
        
        if output:
            chars = len(output.read_text(encoding="utf-8"))
            
            
            print(f"  {source.name} -> {output} ({chars} chars)")

    print("\nDone. Clean text is ready in docs/ for ingestion.")
    
    
if __name__ == "__main__":
    main()